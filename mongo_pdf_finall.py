from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def create_document(data):
    # === Word Document Creation ===
    doc = Document()
    title = doc.add_heading('Video Analysis Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Project Information Section
    doc.add_heading('Project Information', level=1)
    doc.add_paragraph(f"Project ID: {data.get('project_id', 'Not available')}")
    doc.add_paragraph(f"Project Name: {data.get('project_name', 'Not available')}")
    doc.add_paragraph(f"Privacy: {data.get('privacy', 'Not available')}")
    upload_time = data.get('upload_timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    doc.add_paragraph(f"Upload Timestamp: {upload_time}")
    doc.add_paragraph(f"Video Status: {data.get('video_upload_status', 'Not available')}")
    
    # Video Metadata Section
    doc.add_heading('Video Metadata', level=1)
    doc.add_paragraph(f"File Name: {data.get('original_filename', 'Not available')}")
    doc.add_paragraph(f"Video Hash: {data.get('video_hash', 'Not available')}")
    
    metadata = data.get('metadata', {})
    if metadata:
        doc.add_paragraph(f"File Type: {metadata.get('file_type', 'Not available')}")
        doc.add_paragraph(f"File Size: {metadata.get('file_size_MB', 'Not available')} MB")
        doc.add_paragraph(f"Duration: {metadata.get('duration_sec', 'Not available')} seconds")
        doc.add_paragraph(f"Frame Rate: {metadata.get('frame_rate', 'Not available')} fps")
        doc.add_paragraph(f"Resolution: {metadata.get('video_resolution', 'Not available')}")
        
        if 'transcription_with_timestamps' in metadata:
            doc.add_heading('Transcription', level=1)
            for entry in metadata.get('transcription_with_timestamps', []):
                p = doc.add_paragraph()
                p.add_run(f"[{entry.get('start_time', 0):.2f} - {entry.get('end_time', 0):.2f}] ").bold = True
                p.add_run(f"{entry.get('speaker', 'Unknown')}: {entry.get('text', 'No text')}")
        
        if 'summary' in metadata:
            doc.add_heading('Summary', level=1)
            doc.add_paragraph(metadata.get('summary', 'No summary available'))
        
        if 'keywords_and_topics' in metadata:
            doc.add_heading('Keywords and Topics', level=1)
            kt = metadata['keywords_and_topics']
            doc.add_paragraph(f"Keywords: {', '.join(kt.get('keywords', ['None']))}")
            doc.add_paragraph(f"Topics: {', '.join(kt.get('topics', ['None']))}")
        
        if 'brands_and_locations' in metadata:
            doc.add_heading('Brands and Locations', level=1)
            bl = metadata['brands_and_locations']
            doc.add_paragraph(f"Brands: {', '.join(bl.get('brands', ['None']))}")
            doc.add_paragraph(f"Locations: {', '.join(bl.get('locations', ['None']))}")
        
        if 'sentiment' in metadata:
            doc.add_heading('Sentiment', level=1)
            doc.add_paragraph(metadata.get('sentiment', 'Not analyzed'))
        
        if 'people' in metadata:
            doc.add_heading('People', level=1)
            for person in metadata.get('people', []):
                doc.add_paragraph(f"Name: {person.get('name', 'Unknown')}")
                doc.add_paragraph(f"Designation: {person.get('designation', 'Not specified')}")
        
        if 'objects' in metadata:
            doc.add_heading('Objects', level=1)
            objects_list = metadata.get('objects', [])
            objects_text = ", ".join(objects_list) if isinstance(objects_list, list) else "No objects detected"
            doc.add_paragraph(objects_text)
        
        if 'ai_voice_location' in metadata:
            doc.add_heading('AI Voice Location', level=1)
            doc.add_paragraph(metadata.get('ai_voice_location', 'Not available'))

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Generate Word blob
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_blob = word_buffer.getvalue()
    word_buffer.close()
    
    # === PDF Creation ===
    pdf_buffer = BytesIO()
    pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    content = []
    content.append(Paragraph('Video Analysis Report', styles['Title']))
    content.append(Spacer(1, 12))
    
    content.append(Paragraph('Objects', styles['Heading1']))
    objects_text = ", ".join(objects_list) if isinstance(objects_list, list) else "No objects detected"
    content.append(Paragraph(objects_text, styles['Normal']))
    
    # Add footer to PDF
    content.append(Spacer(1, 36))
    content.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    
    # Build PDF
    pdf_doc.build(content)
    pdf_blob = pdf_buffer.getvalue()
    pdf_buffer.close()
    
    return word_blob, pdf_blob

def main(result):
    word_blob, pdf_blob = create_document(result)
    return word_blob, pdf_blob

if __name__ == "__main__":
    from pymongo import MongoClient
    
    def get_video_metadata(video_hash: str):
        client = MongoClient("mongodb+srv://prathameshhh902:Bvit%402002@cluster0.q2fnv.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0")
        db = client['videoDB']
        collection = db['video_processing']
        query = {"video_hash": video_hash}
        video_data = collection.find_one(query)
        client.close()
        return video_data if video_data else None
    
    sample_hash = "ecd71914b736e3664412b10dd127c3447156d2f539181016a90b97cf27fa4e6f"
    result = get_video_metadata(sample_hash)
    if result:
        word_blob, pdf_blob = main(result)
        print(f"Word document blob size: {len(word_blob)} bytes")
        print(f"PDF document blob size: {len(pdf_blob)} bytes")
        
        # Optional: Save blobs to verify contents
        with open("test_output.docx", "wb") as f:
            f.write(word_blob)
        with open("test_output.pdf", "wb") as f:
            f.write(pdf_blob)
    else:
        print("No metadata found for the given video hash")
